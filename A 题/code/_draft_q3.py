# -*- coding: utf-8 -*-
"""生成《问题三_初稿.docx》（风格仿问题一、二初稿；技术细节与 Q3 代码一致）。"""
import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = r"d:\数模校赛\A 题\docs\问题三_初稿.docx"
doc = Document()

# 默认字体
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
from docx.oxml.ns import qn
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def h(text, level=1):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    return p


def para(text, bold=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if align:
        p.alignment = align
    return p


def formula(text):
    """文本式公式段落（居中、斜体可后续在 Word 中替换为公式编辑器对象）。"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


# ============ 标题 ============
para("问题三：分时禁飞区下多无人机巡检的时空调度模型", bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER)

# ============ 1 问题分析 ============
h("1 问题分析", 1)
para("问题三在问题二得到的机队规模与任务副本分配基础上，引入附件 2 给出的分时圆形禁飞区："
     "每个禁飞区具有固定的圆心坐标、半径以及生效时段。全部无人机于 08:00 从基地同时出发，"
     "在禁飞区生效时段内，无人机的任意位置不得落入对应圆形区域的内部。")
para("禁飞区对巡检执行产生两类应对：一是在入界前悬停等待，待管制窗口结束后再沿原直线航段继续飞行；"
     "二是沿禁飞圆外扩安全圆（半径增加 1 个坐标单位，即 100 m）的切线—圆弧折线绕飞。"
     "对每个航段，二者取较早到达者。由于各条路线共享全局时钟，任一路线在某个圆区的等待会推迟其后续"
     "全部活动的时刻，因此本问本质上是连续时间意义下的带时间窗禁行约束的多机重调度问题。")
para("目标仍沿用问题二的层次结构：在固定机队规模下，首先保证总体完成时间尽可能短"
     "（相对无禁飞区基准不再放宽，即 ε=0 主口径），其次压缩各机工作时长极差，"
     "并以方差打破平局。求解时允许调整任务在各路线间的分配与访问顺序，以及每条路线相对 08:00 的循环切入点。")

# ============ 2 模型假设 ============
h("2 模型假设", 1)
para("问题一、问题二中的全部假设继续成立（任务副本、重复巡检计数方式、基地出发与返回、"
     "同机同点任务副本不连续访问、工作时长由飞行时间与服务时间构成），并补充以下约定：")
para("（1）禁飞区生效时段按半开区间 [s_z, e_z) 处理：到达结束时刻后区域立即解除管制，"
     "端点时刻位于边界上视为允许；附件中零长度的时段（起止相同）不构成管制；")
para("（2）无人机可在任意航段的入界点外侧悬停等待，等待时间计入工作时长，但不计入巡检服务时间；")
para("（3）绕飞沿禁飞圆外扩 1 坐标单位（100 m）的安全圆进行，圆弧以不超过 10° 的弦离散，"
     "以保证离散折线仍完全位于原禁飞圆之外；")
para("（4）当同一航段先后经过多个生效圆区或圆区重叠时，等待点取各圆区入界点的共同外侧："
     "沿航段从入界点逐次后退 1 坐标单位探测，直至该点不落入任何生效禁飞圆内部；")
para("（5）若某巡检点位于生效圆区内，其服务时间连同后续在圆内的飞行与任务链统一前视处理，"
     "确保整段圆内占用被整体安排在管制窗口之前或之后，避免在圆内任务链中途非法等待。")

# ============ 3 符号 ============
h("3 新增符号与可行域", 1)
para("令 Z={z_1,…,z_m} 为禁飞区集合，z=(x_z, y_z, r_z, s_z, e_z) 分别表示圆心横纵坐标、"
     "半径（坐标单位）与生效时段的起止时刻（相对 08:00 的秒数）。其余符号定义见表 1。")
para("表 1  问题三新增符号", bold=True)

rows = [
    ("符号", "含义"),
    ("Z, m", "分时禁飞区集合及其个数"),
    ("(x_z, y_z, r_z)", "禁飞区 z 的圆心坐标与半径（坐标单位，1 单位=100 m）"),
    ("(s_z, e_z)", "禁飞区 z 的生效时段（相对 08:00，半开区间）"),
    ("d_i", "无人机 i 的实际飞行距离（含绕行折线）"),
    ("w_i", "无人机 i 的总等待时间"),
    ("T_i", "无人机 i 的动态工作时长（飞行+服务+等待）"),
    ("Tmax, δ", "各机工作时长的最大值与极差（含义同问题二）"),
    ("q", "航段 (a,b) 中位于禁飞圆 z 内的参数区间 [u,v]"),
]
table = doc.add_table(rows=len(rows), cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (a, b) in enumerate(rows):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b

para("")
para("在问题二可行域基础上，可行解还须满足禁飞约束：对任意时刻 t 与任意生效中的禁飞区 z"
     "（s_z ≤ t < e_z），无人机位置 p(t) 满足 ||p(t)−c_z|| ≥ r_z。同一无人机相邻任务副本"
     "不连续的约束仍然保留。")

# ============ 4 模型 ============
h("4 模型", 1)
h("4.1 单航段的连续时间决策", 2)
para("设某航段由点 a 飞向点 b，直线飞行耗时为 Δt，终点服务时间为 s（非巡检点取 0）。"
     "对每个禁飞区 z，先计算线段 ab 落在闭圆内的参数区间 [u,v]（0≤u≤v≤1）：")
formula("u,v = 解方程：||a + q(b−a) − c_z||² = r_z²，q ∈ [0,1]，取交点夹出的区间")
para("出发时刻为 t 时，进入与离开圆区的时刻为")
formula("t_enter = t + w_prev + u·Δt,   t_leave = t + w_prev + v·Δt + s_tail")
para("其中 w_prev 为已累计的等待，s_tail 为终点（含任务链前视）在圆内的占用时长。"
     "若 [t_enter, t_leave) 与管制窗口 [s_z, e_z) 相交，则直飞方案需要在入界前等待")
formula("Δw = e_z − t_enter  （等待点：入界点沿航段后退至所有生效圆外的共同外侧）")
para("绕飞方案由外扩安全圆上两侧切线—圆弧折线构成，取两侧较短者，其飞行距离为 L_detour。"
     "单航段决策取二者较早到达：")
formula("T_leg = min( Δt + Σ_z Δw_z + s ,  L_detour / V + s )")
para("若航段起点已在生效圆内（u=0），则不允许在该航段安排等待，时序必须由前一航段的"
     "任务链前视整体推后（见 4.2）；该情形在评价器中返回不可行。")

h("4.2 圆内任务链前视", 2)
para("当巡检点落入生效圆区时，本点的 300 s 服务与后续连续位于同一圆内的飞行段、"
     "服务构成一段“圆内占用链”。对每个圆区 z 计算自本点起首次离开 z 的时间"
     "clearance_i(z)，将其并入到达本点航段的“终点占用”，从而把整段圆内占用统一安排在"
     "管制窗口之前或之后，避免中途非法等待。")

h("4.3 路线总时长与优化模型", 2)
para("每条路线的动态工作时长 T_i 由其全部航段的 T_leg 累加得到。由于全部无人机同时出发且"
     "共享同一禁飞时间轴，路线相对基地的循环切入点（哪一点先访问）也影响总时长，故对每条路线"
     "还需优化正反向的循环切入位置。优化模型延续问题二的分层结构：")
formula("min  T_max = max_i T_i   （第一层，ε=0：不劣于问题二基准）")
formula("再  min  δ = T_max − T_min  ；平局时取各机工作时长方差较小者")
para("其中等待与绕行决策由 4.1–4.2 的连续时间事件评价器隐式给出，不作为显式变量。")

# ============ 5 求解思路 ============
h("5 求解思路", 1)
para("（1）承接问题二并初始化。 读入问题二的机队规模与路线方案、附件 2 禁飞区数据；"
     "以问题二路线为初始解，构建动态评价器 dynamic_metric。")
para("（2）循环切入优化。 对每条路线穷举正反向全部循环切入点，取动态总时长最短者，"
     "作为后续搜索的起点（cyclic_improve）。")
para("（3）分层遗传算法重调度。 沿用问题二的求解引擎：种群规模 48，锦标赛选择、路径继承交叉、"
     "变邻域变异（路径内 2-opt、路径内 relocate、跨路径 relocate、单任务 swap、成段搬迁、"
     "短任务块交换、循环切入点平移），精英保留；每 10 代对最优个体执行一次“最忙搬点给最闲”"
     "的显式均衡搬迁。适应度按（Tmax，δ，方差）的字典序比较。搜索以问题二路线为初始种群种子，"
     "同时将搜索过程中发现的最好方案继续作为下一轮搜索的起点，多起点重启并保留字典序最优档案，"
     "确保最终采用的方案不劣于搜索过程中得到的最好可行解。")
para("（4）独立几何审计。 用与求解器相独立的审计脚本逐航段复算飞行/等待/服务时长，"
     "并断言：每个物理点累计巡检次数恰为 l_p；同机任务副本不连续；任一活动段的起点、终点"
     "在对应时刻均不位于生效禁飞圆内部；绕飞使用外扩 1 坐标单位的安全圆；等待位置按入界点"
     "外侧极限解释。全部算例审计通过。")
para("（5）结果汇总。 四个算例的最终结果见表 2。")

para("表 2  问题三最终结果（08:00 同时出发；Tmax 与 δ 单位均为小时）", bold=True)
rows2 = [
    ("算例", "N", "Tmax", "δ", "总飞行时长", "总服务时长", "总等待", "绕飞航段数"),
    ("Case1", "4", "8.9260", "0.0041", "29.6951", "6.0000", "0.0000", "0"),
    ("Case2", "2", "11.0692", "0.0927", "7.1790", "11.5833", "3.2832", "4"),
    ("Case3", "5", "9.1271", "0.0468", "33.8616", "11.6667", "0.0000", "1"),
    ("Case4", "4", "9.9006", "0.0375", "24.3795", "15.1667", "0.0000", "0"),
]
table2 = doc.add_table(rows=len(rows2), cols=len(rows2[0]))
table2.style = "Table Grid"
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row in enumerate(rows2):
    for j, val in enumerate(row):
        table2.rows[i].cells[j].text = str(val)
para("注：Case2 的两条路线共发生 4 个绕飞航段与 3 次入界等待，其 Tmax 较问题二基准增加"
     "约 3.43 h；Case3 有 1 个绕飞航段；Case1 与 Case4 无需等待或绕行即可避开全部禁飞窗。")

# ============ 参考文献 ============
h("参考文献", 1)
refs = [
    "[1] MAVROTAS G. Effective implementation of the ε-constraint method in multi-objective mathematical programming problems[J]. Applied Mathematics and Computation, 2009, 213(2): 455–465.",
    "[2] MATL P, HARTL R F, VIDAL T. Workload equity in vehicle routing problems: A survey and analysis[J]. Transportation Science, 2018, 52(2): 239–260.",
    "[3] VIDAL T, CRAINIC T G, GENDREAU M, et al. A hybrid genetic algorithm for multidepot and periodic vehicle routing problems[J]. Operations Research, 2012, 60(3): 611–624.",
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.space_after = Pt(2)

doc.save(OUT)
print("已生成:", OUT)
